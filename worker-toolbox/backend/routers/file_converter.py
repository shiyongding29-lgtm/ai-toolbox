"""File Converter — Word/PDF/Excel/CSV/TXT 互转。"""
import os, tempfile, subprocess
from fastapi import APIRouter, UploadFile, File, Form
from backend.config import config

router = APIRouter(prefix="/api/convert", tags=["file-converter"])


@router.post("/convert")
async def convert_file(file: UploadFile = File(...), to_format: str = Form("txt")):
    contents = await file.read()
    fname = file.filename or 'file'
    src_ext = fname.rsplit('.', 1)[-1].lower() if '.' in fname else ''
    dst_ext = to_format.lower()
    os.makedirs(config.upload_dir, exist_ok=True)
    out_name = f"conv_{os.urandom(4).hex()}.{dst_ext}"
    out_path = os.path.join(config.upload_dir, out_name)

    try:
        if src_ext == 'pdf' and dst_ext in ('txt', 'docx'):
            text = _pdf_to_text(contents)
            if dst_ext == 'docx': _text_to_docx(text, out_path)
            else: open(out_path,'w').write(text)
        elif src_ext == 'docx' and dst_ext == 'txt':
            open(out_path,'w').write(_docx_to_text(contents))
        elif src_ext == 'docx' and dst_ext == 'pdf':
            _libreoffice_convert(contents, 'docx', out_path)
        elif src_ext == 'pptx' and dst_ext == 'pdf':
            _libreoffice_convert(contents, 'pptx', out_path)
        elif src_ext == 'html' and dst_ext == 'pdf':
            _libreoffice_convert(contents, 'html', out_path)
        elif src_ext in ('xlsx','xls') and dst_ext == 'csv':
            import pandas as pd; pd.read_excel(contents).to_csv(out_path, index=False)
        elif src_ext == 'csv' and dst_ext in ('xlsx','xls'):
            import pandas as pd; pd.read_csv(contents).to_excel(out_path, index=False)
        else:
            return {"code":400,"msg":f"Unsupported: {src_ext}→{dst_ext}","data":None}
        return {"code":0,"msg":"ok","data":{"url":f"/uploads/{out_name}","from":src_ext,"to":dst_ext}}
    except Exception as e:
        return {"code":500,"msg":str(e)[:200],"data":None}


def _pdf_to_text(contents):
    import pdfplumber
    with tempfile.NamedTemporaryFile(suffix='.pdf',delete=False) as t: t.write(contents); tp=t.name
    try:
        with pdfplumber.open(tp) as pdf: return '\n\n'.join(p.extract_text() or '' for p in pdf.pages)
    finally: os.unlink(tp)

def _docx_to_text(contents):
    from docx import Document
    with tempfile.NamedTemporaryFile(suffix='.docx',delete=False) as t: t.write(contents); tp=t.name
    try: return '\n'.join(p.text for p in Document(tp).paragraphs)
    finally: os.unlink(tp)

def _text_to_docx(text, out_path):
    from docx import Document
    doc=Document(); [doc.add_paragraph(l) for l in text.split('\n')]; doc.save(out_path)

def _libreoffice_convert(contents, src_ext, out_path):
    """用 LibreOffice 将任意 Office/HTML 文件转为 PDF。"""
    import shutil
    suffix = f'.{src_ext}'
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as t: t.write(contents); tp = t.name
    try:
        subprocess.run(['soffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(out_path), tp],
                       capture_output=True, timeout=60)
        expected = os.path.join(os.path.dirname(out_path), os.path.basename(tp).replace(suffix, '.pdf'))
        if os.path.exists(expected) and expected != out_path:
            shutil.move(expected, out_path)
    finally:
        os.unlink(tp)
